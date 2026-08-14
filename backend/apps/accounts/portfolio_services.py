"""AI-powered portfolio service for the Super Admin.

The Super Admin uploads their own resume; the AI (a) reviews it - pros, cons,
improvements, a score - and (b) builds a public portfolio (headline, about,
skills, education, experience, projects) from it. The AI can also rewrite the
resume into a polished "rebuilt" version, delivered as an editable text
preview plus a downloadable .docx, and then review the rebuilt version too.

Everything here is private to the Super Admin: the resume and both reviews
are never exposed to faculty or students - only the generated portfolio
content (when ``is_published``) is public.

Reuses the placements resume pipeline (text extraction, OCR, provider task
routing) so the portfolio behaves exactly like the student resume analysis.
"""

import io
import re
import secrets
import threading

from django.conf import settings
from django.utils import text as text_utils
from django.utils import timezone
from rest_framework.exceptions import ValidationError

from apps.documents.services import (
    _cloudinary_upload,
    delete_document_file,
    upload_document,
)
from apps.placements.ai import AiError, ai_json, ai_plain_text
from apps.placements.ai_parse import normalize_resume_report
from apps.placements.models import AiUsageLog
from apps.placements.resume_ai import (
    _MAX_TEXT_CHARS,
    _SCANNED_PDF_REASON,
    _deferred_usage,
    _extract_resume_text,
    _is_pdf_file,
    _ocr_resume_pdf,
)

from .models import Portfolio, User

# Portfolio resumes follow the same size target as student resumes (500KB).
PORTFOLIO_RESUME_TARGET_BYTES = 500 * 1024

# Public page look: light/dark mode + an accent color (hex). "auto" follows
# the visitor's own site theme.
DEFAULT_PORTFOLIO_THEME = {"mode": "auto", "accent": "#f56d14"}
_THEME_MODES = ("auto", "light", "dark")


def get_portfolio_theme(portfolio: Portfolio) -> dict:
    """The effective theme with defaults filled in (never raises)."""
    theme = portfolio.theme if isinstance(portfolio.theme, dict) else {}
    mode = theme.get("mode") if theme.get("mode") in _THEME_MODES else "auto"
    accent = str(theme.get("accent") or "").lower()
    if not re.fullmatch(r"#[0-9a-f]{6}", accent):
        accent = DEFAULT_PORTFOLIO_THEME["accent"]
    return {"mode": mode, "accent": accent}


def normalize_portfolio_theme(value) -> dict | None:
    """Validate/clean a theme payload. None => invalid (caller raises a 400)."""
    if not isinstance(value, dict):
        return None
    mode = value.get("mode")
    if mode not in _THEME_MODES:
        return None
    accent = str(value.get("accent") or "").lower()
    if not re.fullmatch(r"#[0-9a-f]{6}", accent):
        return None
    return {"mode": mode, "accent": accent}

# One AI call does both jobs: the private review AND the public portfolio
# content. Keeps the analysis cheap (the resume text is sent once).
_PORTFOLIO_REVIEW_PROMPT = """\
You are a career advisor building a personal portfolio for a college Super Admin / placement head.
Their resume text is provided. Return ONLY a single valid JSON object - no markdown, no code fences, no prose before or after.
The object must use EXACTLY this schema:
{
  "score": 0-100 integer (resume quality),
  "summary": "two or three sentences on the resume's overall impression",
  "pros": ["4-6 short strings - what the resume genuinely does WELL"],
  "cons": ["3-5 short strings - genuine weaknesses or risks"],
  "improvements": ["5-8 COMPLETE, concrete action items - one short sentence each on exactly what to add, fix or quantify, ordered by impact"],
  "skills": ["every skill/keyword mentioned - languages, tools, frameworks, soft skills"],
  "ats_keywords": ["important ATS keywords for IT/management roles MISSING from this resume"],
  "headline": "one short professional headline for their public portfolio (e.g. 'Placement Head & Software Engineer')",
  "about": "2-3 sentence professional bio written in the first person for their portfolio",
  "education": "one short paragraph describing their education, or \"\" if none",
  "experience": "one short paragraph summarising their work/leadership experience, or \"\" if none",
  "projects": "one short paragraph summarising their notable projects, or \"\" if none"
}
Be specific and honest. Never invent experience, education or projects that are not in the resume -
use \"\" for sections the resume does not mention. If the resume text is unreadable or empty, still return
the JSON with score 0 and a note in summary that the text could not be extracted.
"""

# When the owner pastes their ORIGINAL resume source code (e.g. LaTeX), the
# rebuild keeps that exact structure and only strengthens the content.
_SOURCE_FILL_PROMPT = """\
You are improving a resume. Below is the user's ORIGINAL resume source code (LaTeX or plain text)
followed by improvement notes about their content.
Rewrite ONLY the content - strengthen summaries, bullets, skill lists and wording - while keeping
EVERY structural element byte-for-byte identical: all LaTeX commands, packages, section headings,
formatting, spacing and the overall layout must remain untouched. Do NOT add or remove sections,
and do NOT change any \\section / \\subsection / documentclass / usepackage lines.
Return ONLY the complete updated source code in a single ```latex code block (or a plain ``` block
if the source is not LaTeX). Preserve everything else exactly.
"""

# The AI rewrite: structured sections for the polished ("rebuilt") resume.
_REBUILD_PROMPT = """\
You are a professional resume writer. Rewrite the provided resume into a polished, ATS-friendly version.
Keep every fact (names, dates, roles, projects, skills) EXACTLY as stated - improve wording, structure
and impact, never invent new facts. Return ONLY a single valid JSON object - no markdown, no code fences:
{
  "summary": "2-3 sentence professional summary with strong action verbs",
  "skills": ["10-15 skills grouped sensibly - technical and soft"],
  "experience": "rewritten experience section as plain text - bullets separated by newlines, each starting with '- '",
  "projects": "rewritten projects section as plain text - bullets separated by newlines, each starting with '- '",
  "education": "one short paragraph with the education details, or \"\" if none"
}
If a section is not in the original resume, use \"\" (or [] for skills). Never invent facts.
"""


def generate_portfolio_slug(user: User) -> str:
    """A unique public link slug for a portfolio (e.g. ``admin123-4f2a9c``)."""
    base = text_utils.slugify(user.roll_number or "admin")[:30] or "admin"
    for _ in range(20):
        slug = f"{base}-{secrets.token_hex(3)}"
        if not Portfolio.objects.filter(slug=slug).exists():
            return slug
    return f"{base}-{secrets.token_hex(5)}"


def portfolio_folder(portfolio: Portfolio) -> str:
    """Cloudinary folder: ``portfolios/{roll}/`` (mirrors the resume layout)."""
    return f"portfolios/{text_utils.slugify(portfolio.user.roll_number) or 'admin'}"


def sync_portfolio_resume_ref(portfolio: Portfolio, student: User) -> bool:
    """Point a student's portfolio at their latest Resume file (no AI involved).

    A student's portfolio is built from the resume they already uploaded for
    faculty - there is no separate portfolio upload. This copies the file
    references over so the analysis pipeline (extraction/OCR/review) reads the
    right file. Returns True when a usable resume exists, False when the
    student has no resume yet. AI fields are left untouched.
    """
    from .models import Resume

    resume = (
        Resume.objects.filter(student=student)
        .exclude(is_missing=True)
        .order_by("-updated_at")
        .first()
    )
    if not resume or not resume.public_id:
        return False
    portfolio.file_name = resume.file_name
    portfolio.file_size = resume.file_size
    portfolio.cloudinary_url = resume.cloudinary_url
    portfolio.public_id = resume.public_id
    portfolio.is_missing = False
    portfolio.save(update_fields=[
        "file_name", "file_size", "cloudinary_url", "public_id", "is_missing",
        "updated_at",
    ])
    return True


def _as_number(value, default: float, lo: float, hi: float) -> float:
    """Coerce a number, clamped into [lo, hi], falling back to ``default``."""
    try:
        num = float(value)
    except (TypeError, ValueError):
        return default
    return max(lo, min(hi, num))


def clean_portfolio_images(value) -> list | None:
    """Validate/normalize the placed-images list. None => invalid."""
    if not isinstance(value, list):
        return None
    out: list[dict] = []
    for item in value:
        if not isinstance(item, dict):
            return None
        url = str(item.get("url") or "").strip()
        if not url or len(url) > 500:
            return None
        out.append({
            "url": url,
            "public_id": str(item.get("public_id") or "")[:255],
            "alt": str(item.get("alt") or "")[:150],
            "x": _as_number(item.get("x"), 50, 0, 100),
            "y": _as_number(item.get("y"), 50, 0, 100),
            "width": _as_number(item.get("width"), 200, 40, 900),
            "height": _as_number(item.get("height"), 200, 40, 900),
            "opacity": _as_number(item.get("opacity"), 1, 0, 1),
        })
        if len(out) >= 12:
            break
    return out


def clean_background_image(value) -> dict | None:
    """Validate/normalize the background image. None => invalid. Empty => None."""
    if value is None or value == {} or value == [] or value == "":
        return None
    if not isinstance(value, dict):
        return None
    url = str(value.get("url") or "").strip()
    if not url or len(url) > 500:
        return None
    return {
        "url": url,
        "public_id": str(value.get("public_id") or "")[:255],
        "opacity": _as_number(value.get("opacity"), 0.35, 0, 1),
        "darken": _as_number(value.get("darken"), 0.55, 0, 1),
    }


_IMAGE_TYPES = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
_IMAGE_MAX_BYTES = 8 * 1024 * 1024


def _validate_portfolio_image(image_file) -> None:
    """Lightweight image check (extension + size)."""
    if not image_file or image_file.size <= 0:
        raise ValidationError({"file": "An image file is required."})
    if image_file.size > _IMAGE_MAX_BYTES:
        raise ValidationError({
            "file": f"The image is larger than {_IMAGE_MAX_BYTES // (1024 * 1024)}MB."
        })
    name = (getattr(image_file, "name", "") or "").lower()
    _, _, ext_part = name.rpartition(".")
    ext = f".{ext_part}" if ext_part else ""
    if ext not in _IMAGE_TYPES:
        raise ValidationError({"file": "Only JPG, PNG, GIF or WebP images are allowed."})


def upload_portfolio_image(portfolio: Portfolio, image_file, request=None) -> dict:
    """Upload a portfolio image to Cloudinary and return its references."""
    from apps.core.utils import log_audit

    _validate_portfolio_image(image_file)
    try:
        result = _cloudinary_upload(
            image_file, f"{portfolio_folder(portfolio)}/images", resource_type="image"
        )
    except Exception as exc:
        raise ValidationError({"file": f"Cloudinary upload failed: {exc}"})
    log_audit(
        portfolio.user, "PORTFOLIO_IMAGE_UPLOAD", "Portfolio", portfolio.id,
        {"roll_number": portfolio.user.roll_number, "file": image_file.name},
        request,
    )
    return {
        "url": result["secure_url"],
        "public_id": result["public_id"],
        "file_name": image_file.name,
    }


def delete_portfolio_image(public_id: str) -> bool:
    """Remove a portfolio image from Cloudinary (best-effort)."""
    if not public_id:
        return False
    return delete_document_file(public_id)


def _reset_analysis(portfolio: Portfolio, *, rebuilt: bool = True) -> None:
    """Clear every AI/cache field (used when a new file replaces the old one).

    ``resume_source`` (the owner's own template) is left untouched - it is
    their input, not a generated field.
    """
    portfolio.ai_status = Portfolio.AiStatus.PENDING
    portfolio.ai_score = None
    portfolio.ai_analysis = None
    portfolio.ai_error = ""
    portfolio.ai_analyzed_at = None
    # The content was generated from the old file - it no longer applies.
    portfolio.headline = ""
    portfolio.about = ""
    portfolio.skills = []
    portfolio.education = ""
    portfolio.experience = ""
    portfolio.projects = ""
    if rebuilt:
        portfolio.rebuilt_sections = None
        portfolio.rebuilt_text = ""
        portfolio.rebuilt_file_name = ""
        portfolio.rebuilt_docx_url = ""
        portfolio.rebuilt_docx_public_id = ""
        portfolio.rebuilt_tex = ""
        portfolio.rebuilt_pdf_url = ""
        portfolio.rebuilt_pdf_public_id = ""
        portfolio.rebuilt_at = None
        portfolio.rebuilt_ai_status = Portfolio.AiStatus.PENDING
        portfolio.rebuilt_ai_score = None
        portfolio.rebuilt_ai_analysis = None
        portfolio.rebuilt_ai_error = ""
        portfolio.rebuilt_ai_analyzed_at = None


def upload_portfolio_resume(portfolio: Portfolio, resume_file, request=None) -> Portfolio:
    """Upload/replace the admin's portfolio resume on Cloudinary.

    Resets all previous analysis - the new file gets fresh reviews. Returns
    the portfolio (the caller serializes it). Never raises for AI issues;
    Cloudinary/validation errors bubble up as DRF ValidationErrors.
    """
    from apps.core.utils import log_audit

    uploaded = upload_document(
        resume_file, portfolio_folder(portfolio),
        target_bytes=PORTFOLIO_RESUME_TARGET_BYTES,
    )
    created = not portfolio.public_id
    if not created and portfolio.public_id != uploaded["public_id"]:
        delete_document_file(portfolio.public_id)

    portfolio.file_name = uploaded["file_name"]
    portfolio.file_size = uploaded["file_size"]
    portfolio.cloudinary_url = uploaded["url"]
    portfolio.public_id = uploaded["public_id"]
    portfolio.is_missing = False
    _reset_analysis(portfolio)
    portfolio.save()
    log_audit(
        portfolio.user, "PORTFOLIO_RESUME_UPLOAD" if created else "PORTFOLIO_RESUME_UPDATE",
        "Portfolio", portfolio.id,
        {"roll_number": portfolio.user.roll_number, "file": portfolio.file_name},
        request,
    )
    # Auto-analyse in the background so the review + portfolio content are
    # ready when the admin opens the builder. Gated like student resumes.
    if getattr(settings, "AI_AUTO_ANALYZE_ON_UPLOAD", True):
        try:
            threading.Thread(
                target=_auto_analyze_in_thread, args=(portfolio.id,), daemon=True
            ).start()
        except Exception:
            pass  # never fail an upload because the background thread failed
    return portfolio


def _auto_analyze_in_thread(portfolio_id: int) -> None:
    from django.db import close_old_connections

    close_old_connections()
    try:
        portfolio = Portfolio.objects.select_related("user").filter(pk=portfolio_id).first()
        if not portfolio or portfolio.is_missing or not portfolio.public_id:
            return
        analyze_portfolio(portfolio, portfolio.user)
    except Exception:
        pass  # background analysis must never crash anything
    finally:
        close_old_connections()


def _normalize_portfolio_report(raw) -> dict | None:
    """Normalize the review+content response into a single dict.

    Reuses the resume-report normalizer for the review fields and adds the
    portfolio content fields with sensible fallbacks (headline -> user's name,
    about -> the review summary, missing sections -> ""). Returns None when
    the response carried no usable review at all.
    """
    review = normalize_resume_report(raw)
    if review is None:
        return None
    if not isinstance(raw, dict):
        raw = {}
    # Unwrap wrapper keys the same way normalize_resume_report does.
    for wrapper in ("report", "result", "data", "analysis", "quality"):
        if isinstance(raw.get(wrapper), dict) and len(raw) == 1:
            raw = raw[wrapper]
            break
    if not isinstance(raw, dict):
        raw = {}

    def _txt(*names: str) -> str:
        for name in names:
            value = raw.get(name)
            if value is not None and str(value).strip():
                return str(value).strip()
        return ""

    return {
        **review,
        "headline": _txt("headline", "title", "tagline"),
        "about": _txt("about", "bio", "profile"),
        "education": _txt("education", "education_details"),
        "experience": _txt("experience", "work_experience", "work"),
        "projects": _txt("projects", "project_experience"),
    }


def analyze_portfolio(portfolio: Portfolio, actor: User) -> dict:
    """Run the AI review + portfolio generation and store it on the portfolio.

    Raises AiError when the AI service itself fails (the caller surfaces a
    502). Credits are only committed when the analysis completes - a failed
    or unreadable run never burns the AI budget.
    """
    text, read_error = _extract_resume_text(portfolio)
    usage, commit_usage = _deferred_usage(actor)
    ocr_used = False

    if not text or not text.strip():
        if read_error == _SCANNED_PDF_REASON and _is_pdf_file(portfolio):
            ocr_text = _ocr_resume_pdf(portfolio, usage)
            if ocr_text and ocr_text.strip():
                text = ocr_text.strip()
                read_error = ""
                ocr_used = True

    if not text or not text.strip():
        portfolio.ai_status = Portfolio.AiStatus.FAILED
        portfolio.ai_score = None
        portfolio.ai_analysis = None
        portfolio.ai_error = (
            "The AI could not read this resume - " + read_error
            if read_error
            else "The AI could not read this resume. Try a text-based PDF."
        )
        portfolio.ai_analyzed_at = timezone.now()
        portfolio.save(update_fields=[
            "ai_status", "ai_score", "ai_analysis", "ai_error",
            "ai_analyzed_at", "updated_at",
        ])
        return _portfolio_payload(portfolio)

    brief = text.strip()[:_MAX_TEXT_CHARS]
    try:
        raw = ai_json(
            _PORTFOLIO_REVIEW_PROMPT, brief, max_tokens=4000,
            usage_callback=usage, task="RESUME_ANALYSIS",
        )
    except AiError:
        raise  # nothing collected, nothing committed - no credits burned

    report = _normalize_portfolio_report(raw)
    if report is None:
        portfolio.ai_status = Portfolio.AiStatus.FAILED
        portfolio.ai_score = None
        portfolio.ai_analysis = None
        portfolio.ai_error = (
            "The AI service returned an unreadable report - the provider "
            "answered, but its model output couldn't be read as a report. "
            "No credits were used for this attempt."
        )
        portfolio.ai_analyzed_at = timezone.now()
        portfolio.save(update_fields=[
            "ai_status", "ai_score", "ai_analysis", "ai_error",
            "ai_analyzed_at", "updated_at",
        ])
        return _portfolio_payload(portfolio)

    if ocr_used:
        report["ocr"] = True
    portfolio.ai_status = Portfolio.AiStatus.COMPLETE
    portfolio.ai_score = report["score"]
    portfolio.ai_analysis = report
    portfolio.ai_error = ""
    portfolio.ai_analyzed_at = timezone.now()
    # Public portfolio content - auto-built from the resume (editable later).
    user = portfolio.user
    portfolio.headline = report["headline"] or (user.full_name or "").strip()[:200]
    portfolio.about = report["about"] or report["summary"]
    portfolio.skills = report["skills"]
    portfolio.education = report["education"]
    portfolio.experience = report["experience"]
    portfolio.projects = report["projects"]
    commit_usage()
    portfolio.save()
    return _portfolio_payload(portfolio)


def _normalize_rebuilt_sections(raw) -> dict | None:
    """Coerce the rebuild response into {summary, skills[], experience, projects, education}."""
    if not isinstance(raw, dict) or not raw:
        return None
    for wrapper in ("report", "result", "data", "resume", "rewritten"):
        if isinstance(raw.get(wrapper), dict) and len(raw) == 1:
            raw = raw[wrapper]
            break
    if not isinstance(raw, dict):
        return None

    def _txt(*names: str) -> str:
        for name in names:
            value = raw.get(name)
            if value is not None and str(value).strip():
                return str(value).strip()
        return ""

    skills_raw = raw.get("skills") or raw.get("skill_set") or []
    if isinstance(skills_raw, str):
        skills = [s.strip().lstrip("-*• ") for s in skills_raw.replace(";", ",").split(",") if s.strip()]
    elif isinstance(skills_raw, list):
        skills = [str(s).strip() for s in skills_raw if str(s).strip()]
    else:
        skills = []

    summary = _txt("summary", "professional_summary", "objective")
    experience = _txt("experience", "work_experience", "work")
    projects = _txt("projects", "project_experience")
    education = _txt("education", "education_details")
    if not summary and not skills and not experience and not projects and not education:
        return None
    return {
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
    }


def _sections_to_text(sections: dict) -> str:
    """Plain-text rendering of the rebuilt sections (for the review + preview)."""
    lines: list[str] = []
    if sections.get("summary"):
        lines.append("PROFESSIONAL SUMMARY\n" + sections["summary"])
    if sections.get("skills"):
        lines.append("\nSKILLS\n" + ", ".join(sections["skills"]))
    for key, title in (("experience", "EXPERIENCE"), ("projects", "PROJECTS"), ("education", "EDUCATION")):
        value = sections.get(key)
        if value:
            lines.append(f"\n{title}\n{value}")
    return "\n\n".join(lines)


def _build_resume_docx(user: User, headline: str, sections: dict) -> bytes:
    """Render the rebuilt resume as a .docx (python-docx) and return its bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    # Base font: clean, professional look.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(user.full_name or user.roll_number)
    run.bold = True
    run.font.size = Pt(22)

    contact_parts = [part for part in (
        user.roll_number,
        user.email or "",
        user.phone or "",
    ) if part]
    if contact_parts or headline:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(" | ".join(contact_parts))
        if headline:
            head = doc.add_paragraph()
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head.add_run(headline).italic = True

    def section(title_text: str, body: str) -> None:
        if not body.strip():
            return
        h = doc.add_heading(title_text, level=1)
        for run_ in h.runs:
            run_.font.size = Pt(13)
        for line in body.splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(line.strip().lstrip("-*• ").strip())

    section("Professional Summary", sections.get("summary", ""))
    if sections.get("skills"):
        h = doc.add_heading("Skills", level=1)
        for run_ in h.runs:
            run_.font.size = Pt(13)
        doc.add_paragraph(", ".join(sections["skills"]))
    section("Experience", sections.get("experience", ""))
    section("Projects", sections.get("projects", ""))
    section("Education", sections.get("education", ""))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


_LATEX_ESCAPE = str.maketrans({
    "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
    "_": r"\_", "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
})


def _sections_to_latex(user: User, headline: str, sections: dict) -> str:
    """Generate a clean LaTeX resume from the rebuilt sections."""
    esc = lambda t: str(t).translate(_LATEX_ESCAPE)
    name = esc(user.full_name or user.roll_number)
    contact = " | ".join(part for part in (user.roll_number, user.email or "", user.phone or "") if part)
    contact = esc(contact)
    head = esc(headline)
    lines = [
        r"% " + name + r" - rebuilt by PlaceMate AI",
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[margin=0.7in]{geometry}",
        r"\usepackage{parskip}",
        r"\usepackage{enumitem}",
        r"\usepackage{hyperref}",
        r"\hypersetup{hidelinks}",
        r"\begin{document}",
        r"\begin{center}",
        r"  {\LARGE\bfseries " + name + r"}\\[2pt]",
    ]
    if head:
        lines.append(r"  {\large\itshape " + head + r"}\\[2pt]")
    if contact:
        lines.append(r"  {\small " + contact + r"}")
    lines.append(r"\end{center}")

    def add_section(title: str, body: str) -> None:
        if not body.strip():
            return
        lines.append(r"\section*{" + esc(title) + r"}")
        for bullet in body.splitlines():
            text = bullet.strip().lstrip("-*• ").strip()
            if text:
                lines.append(r"\begin{itemize}\item " + esc(text) + r"\end{itemize}")

    lines.append(r"\section*{Professional Summary}")
    lines.append(esc(sections.get("summary", "")))
    if sections.get("skills"):
        lines.append(r"\section*{Skills}")
        lines.append(esc(", ".join(sections["skills"])))
    add_section("Experience", sections.get("experience", ""))
    add_section("Projects", sections.get("projects", ""))
    add_section("Education", sections.get("education", ""))
    lines.append(r"\end{document}")
    return "\n".join(lines)


def _strip_latex_fences(text: str) -> str:
    """Remove ```latex / ``` wrappers the model may add around the source."""
    import re as _re

    text = _re.sub(r"```(?:latex|tex)?\s*\n?", "", text, flags=_re.IGNORECASE)
    return text.replace("```", "").strip()


def _build_resume_pdf(user: User, headline: str, sections: dict) -> bytes:
    """Render the rebuilt resume as a clean PDF (reportlab) and return bytes."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
        title=f"{user.full_name or user.roll_number} - Resume",
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("Title", parent=base["Title"], alignment=TA_CENTER, fontSize=22, spaceAfter=2),
        "head": ParagraphStyle("Head", parent=base["Heading2"], textColor=colors.HexColor("#f56d14"), fontSize=12, spaceBefore=10, spaceAfter=4),
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontSize=10, leading=14),
    }
    story = [Paragraph(_xml_escape(user.full_name or user.roll_number), styles["title"])]
    if headline:
        story.append(Paragraph(_xml_escape(headline), ParagraphStyle("Tag", parent=styles["body"], alignment=TA_CENTER, italic=True, spaceAfter=6)))
    contact = " | ".join(p for p in (user.roll_number, user.email or "", user.phone or "") if p)
    if contact:
        story.append(Paragraph(_xml_escape(contact), ParagraphStyle("Contact", parent=styles["body"], alignment=TA_CENTER, spaceAfter=8)))

    def add_section(title: str, body: str) -> None:
        if not body.strip():
            return
        story.append(Paragraph(_xml_escape(title), styles["head"]))
        for line in body.splitlines():
            text = line.strip().lstrip("-*• ").strip()
            if text:
                story.append(Paragraph("&bull;&nbsp;" + _xml_escape(text), styles["body"]))

    story.append(Paragraph("Professional Summary", styles["head"]))
    story.append(Paragraph(_xml_escape(sections.get("summary", "")), styles["body"]))
    if sections.get("skills"):
        story.append(Paragraph("Skills", styles["head"]))
        story.append(Paragraph(_xml_escape(", ".join(sections["skills"])), styles["body"]))
    add_section("Experience", sections.get("experience", ""))
    add_section("Projects", sections.get("projects", ""))
    add_section("Education", sections.get("education", ""))
    doc.build(story)
    return buffer.getvalue()


def _xml_escape(text: str) -> str:
    import html

    return html.escape(str(text), quote=True)


def rebuild_resume(portfolio: Portfolio, actor: User) -> dict:
    """AI-rewrite the resume into a polished version, then review that version.

    Generates a downloadable .docx (stored on Cloudinary), keeps the editable
    plain-text preview, and runs a fresh review on the rebuilt text so the
    admin sees the improvement ("the again review"). Credits are committed
    only when the whole run completes; a failure charges nothing.
    """
    from django.core.files.uploadedfile import SimpleUploadedFile

    from apps.core.utils import log_audit

    text, read_error = _extract_resume_text(portfolio)
    usage, commit_usage = _deferred_usage(actor)

    if not text or not text.strip():
        if read_error == _SCANNED_PDF_REASON and _is_pdf_file(portfolio):
            ocr_text = _ocr_resume_pdf(portfolio, usage)
            if ocr_text and ocr_text.strip():
                text = ocr_text.strip()
    if not text or not text.strip():
        portfolio.rebuilt_ai_status = Portfolio.AiStatus.FAILED
        portfolio.rebuilt_ai_error = (
            "The AI could not read this resume - " + read_error
            if read_error
            else "The AI could not read this resume. Try a text-based PDF."
        )
        portfolio.rebuilt_ai_score = None
        portfolio.rebuilt_ai_analysis = None
        portfolio.rebuilt_ai_analyzed_at = timezone.now()
        portfolio.save(update_fields=[
            "rebuilt_ai_status", "rebuilt_ai_error", "rebuilt_ai_score",
            "rebuilt_ai_analysis", "rebuilt_ai_analyzed_at", "updated_at",
        ])
        return _portfolio_payload(portfolio)

    brief = text.strip()[:_MAX_TEXT_CHARS]
    try:
        raw = ai_json(
            _REBUILD_PROMPT, brief, max_tokens=4000,
            usage_callback=usage, task="RESUME_ANALYSIS",
        )
    except AiError:
        raise  # nothing committed - no credits burned

    sections = _normalize_rebuilt_sections(raw)
    if sections is None:
        portfolio.rebuilt_ai_status = Portfolio.AiStatus.FAILED
        portfolio.rebuilt_ai_score = None
        portfolio.rebuilt_ai_analysis = None
        portfolio.rebuilt_ai_error = (
            "The AI service returned an unreadable report - the provider "
            "answered, but its model output couldn't be read as a report. "
            "No credits were used for this attempt."
        )
        portfolio.rebuilt_ai_analyzed_at = timezone.now()
        portfolio.save(update_fields=[
            "rebuilt_ai_status", "rebuilt_ai_score", "rebuilt_ai_analysis",
            "rebuilt_ai_error", "rebuilt_ai_analyzed_at", "updated_at",
        ])
        return _portfolio_payload(portfolio)

    rebuilt_text = _sections_to_text(sections)

    # LaTeX: when the owner pasted their original source code, fill that EXACT
    # template (keeps their original layout - only the content is improved);
    # otherwise generate a clean LaTeX resume from the improved sections.
    rebuilt_tex = ""
    if (portfolio.resume_source or "").strip():
        try:
            filled = ai_plain_text(
                _SOURCE_FILL_PROMPT,
                (
                    portfolio.resume_source
                    + "\n\nIMPROVEMENT NOTES FROM YOUR ORIGINAL RESUME:\n"
                    + brief
                )[:_MAX_TEXT_CHARS],
                max_tokens=8000, usage_callback=usage, task="RESUME_ANALYSIS",
            )
            if filled and filled.strip():
                rebuilt_tex = _strip_latex_fences(filled)[:20000]
        except AiError:
            rebuilt_tex = ""  # fall back to the generated LaTeX below
    if not rebuilt_tex:
        rebuilt_tex = _sections_to_latex(portfolio.user, portfolio.headline, sections)

    # Review the rebuilt version - a second, cheaper call on the new text.
    rebuilt_review = None
    rebuilt_score = None
    try:
        raw_review = ai_json(
            _PORTFOLIO_REVIEW_PROMPT, rebuilt_text[:_MAX_TEXT_CHARS],
            max_tokens=3500, usage_callback=usage, task="RESUME_ANALYSIS",
        )
        rebuilt_review = normalize_resume_report(raw_review)
    except AiError:
        # The rebuild itself succeeded - a review failure should not discard
        # the polished resume; it just leaves the review PENDING/empty.
        rebuilt_review = None
    if rebuilt_review:
        rebuilt_score = rebuilt_review["score"]

    # Build + upload the .docx (best-effort - the text preview always works).
    docx_bytes = None
    try:
        docx_bytes = _build_resume_docx(portfolio.user, portfolio.headline, sections)
    except Exception:
        docx_bytes = None
    old_public_id = portfolio.rebuilt_docx_public_id
    new_url = ""
    new_public_id = ""
    new_file_name = ""
    if docx_bytes:
        try:
            uploaded = upload_document(
                SimpleUploadedFile(
                    f"{text_utils.slugify(portfolio.user.full_name or portfolio.user.roll_number) or 'resume'}-rebuilt.docx",
                    docx_bytes,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                portfolio_folder(portfolio),
                target_bytes=PORTFOLIO_RESUME_TARGET_BYTES,
            )
            new_url = uploaded["url"]
            new_public_id = uploaded["public_id"]
            new_file_name = uploaded["file_name"]
            if old_public_id and old_public_id != new_public_id:
                delete_document_file(old_public_id)
        except Exception:
            new_url = ""
            new_public_id = ""
            new_file_name = ""

    # Render + upload the .pdf (best-effort - a clean structured PDF).
    pdf_bytes = None
    try:
        pdf_bytes = _build_resume_pdf(portfolio.user, portfolio.headline, sections)
    except Exception:
        pdf_bytes = None
    old_pdf_public_id = portfolio.rebuilt_pdf_public_id
    new_pdf_url = ""
    new_pdf_public_id = ""
    if pdf_bytes:
        try:
            pdf_uploaded = upload_document(
                SimpleUploadedFile(
                    f"{text_utils.slugify(portfolio.user.full_name or portfolio.user.roll_number) or 'resume'}-rebuilt.pdf",
                    pdf_bytes,
                    content_type="application/pdf",
                ),
                portfolio_folder(portfolio),
                target_bytes=PORTFOLIO_RESUME_TARGET_BYTES,
            )
            new_pdf_url = pdf_uploaded["url"]
            new_pdf_public_id = pdf_uploaded["public_id"]
            if old_pdf_public_id and old_pdf_public_id != new_pdf_public_id:
                delete_document_file(old_pdf_public_id)
        except Exception:
            new_pdf_url = ""
            new_pdf_public_id = ""

    portfolio.rebuilt_sections = sections
    portfolio.rebuilt_text = rebuilt_text
    portfolio.rebuilt_file_name = new_file_name
    portfolio.rebuilt_docx_url = new_url
    portfolio.rebuilt_docx_public_id = new_public_id
    portfolio.rebuilt_tex = rebuilt_tex
    portfolio.rebuilt_pdf_url = new_pdf_url
    portfolio.rebuilt_pdf_public_id = new_pdf_public_id
    portfolio.rebuilt_at = timezone.now()
    portfolio.rebuilt_ai_status = (
        Portfolio.AiStatus.COMPLETE if rebuilt_review else Portfolio.AiStatus.PENDING
    )
    portfolio.rebuilt_ai_score = rebuilt_score
    portfolio.rebuilt_ai_analysis = rebuilt_review
    portfolio.rebuilt_ai_error = "" if rebuilt_review else (
        "The review of the rebuilt resume could not be generated - the rebuild itself succeeded."
    )
    portfolio.rebuilt_ai_analyzed_at = timezone.now()
    commit_usage()
    portfolio.save()
    log_audit(
        actor, "PORTFOLIO_REBUILD", "Portfolio", portfolio.id,
        {"roll_number": portfolio.user.roll_number}, None,
    )
    return _portfolio_payload(portfolio)


def delete_portfolio_resume(portfolio: Portfolio, actor: User, request=None) -> None:
    """Remove the portfolio resume (and the rebuilt .docx) from Cloudinary."""
    from apps.core.utils import log_audit

    roll = portfolio.user.roll_number
    if portfolio.public_id:
        delete_document_file(portfolio.public_id)
    if portfolio.rebuilt_docx_public_id:
        delete_document_file(portfolio.rebuilt_docx_public_id)
    if portfolio.rebuilt_pdf_public_id:
        delete_document_file(portfolio.rebuilt_pdf_public_id)
    portfolio.file_name = ""
    portfolio.file_size = 0
    portfolio.cloudinary_url = ""
    portfolio.public_id = ""
    portfolio.is_missing = False
    _reset_analysis(portfolio)
    portfolio.save()
    log_audit(actor, "PORTFOLIO_RESUME_DELETE", "Portfolio", portfolio.id,
              {"roll_number": roll}, request)


def _portfolio_payload(portfolio: Portfolio) -> dict:
    """Small dict of the fields the views return after an action."""
    return {
        "ai_status": portfolio.ai_status,
        "ai_score": portfolio.ai_score,
        "ai_analysis": portfolio.ai_analysis,
        "ai_error": portfolio.ai_error,
        "rebuilt_ai_status": portfolio.rebuilt_ai_status,
        "rebuilt_ai_score": portfolio.rebuilt_ai_score,
        "rebuilt_ai_analysis": portfolio.rebuilt_ai_analysis,
        "rebuilt_ai_error": portfolio.rebuilt_ai_error,
    }
